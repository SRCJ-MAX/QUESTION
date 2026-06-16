from __future__ import annotations

import html
import json
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document


OPTION_RE = re.compile(r"^([A-F])[\s\.、．]?(.*)$", re.IGNORECASE)
NUMBER_RE = re.compile(r"^\s*(\d+)[\.、．]\s*(.*)$")
ANSWER_ITEM_RE = re.compile(r"(\d+)\s*[\.、]?\s*([A-F]+)", re.IGNORECASE)
RANGE_ANSWER_RE = re.compile(r"(\d+)\s*[-－]\s*(\d+)\s*([A-F]+)", re.IGNORECASE)
AI_TYPE_RE = re.compile(r"【[^】]*(单选题|多选题|判断题)[^】]*】")


@dataclass
class ParsedQuestion:
    type: str
    question: str
    options: list[str]
    answer: list[str]
    analysis: str
    chapter: str


def clean_text(value: str) -> str:
    value = html.unescape(value).replace("\xa0", " ")
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def split_embedded_markers(lines: list[str]) -> list[str]:
    result: list[str] = []
    for line in lines:
        text = clean_text(line)
        text = re.sub(r"\s+(?=[（(][一二三四五六七八九十][）)]\s*(?:单选题|多选题|判断题))", "\n", text)
        text = re.sub(r"\s+(?=二[、\.．]\s*简答题)", "\n", text)
        text = re.sub(r"\s+(?=\d{1,3}[\.、．])", "\n", text)
        text = re.sub(r"\s+(?=[A-Fa-f][\.、．])", "\n", text)
        text = re.sub(r"\s+(?=[A-Fa-f][\u4e00-\u9fff0-9])", "\n", text)
        text = re.sub(r"(\S)(参考答案)", r"\1\n\2", text)
        for part in text.splitlines():
            part = clean_text(part)
            if part:
                result.append(part)
    return result


def read_docx(path: Path) -> list[str]:
    doc = Document(path)
    lines: list[str] = []
    for para in doc.paragraphs:
        text = clean_text(para.text)
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            text = clean_text(" ".join(cell.text for cell in row.cells))
            if text:
                lines.append(text)
    return lines


def read_html_doc(path: Path) -> list[str]:
    raw = path.read_bytes().decode("utf-16le", errors="ignore")
    body_start = raw.find("<body")
    body = raw[body_start:] if body_start >= 0 else raw
    body = re.sub(r"(?is)<style.*?</style>|<xml.*?</xml>|<!--.*?-->", "\n", body)
    body = re.sub(r"(?i)<br\s*/?>|</p>|</div>|</tr>|</h\d>", "\n", body)
    body = re.sub(r"(?is)<[^>]+>", "", body)
    return [clean_text(line) for line in body.splitlines() if clean_text(line)]


def read_binary_doc_runs(path: Path) -> list[str]:
    raw = path.read_bytes().decode("utf-16le", errors="ignore")
    runs = re.findall(r"[\u4e00-\u9fffA-Za-z0-9，。、“”‘’（）()；;：:？！、\s\r\n\.\-—_\xa0]{2,}", raw)
    lines: list[str] = []
    for run in runs:
        for part in re.split(r"[\r\n]+", run):
            text = clean_text(part)
            if not text:
                continue
            if any(noise in text for noise in ["WPS Office", "攑h", "鴀錀", "耀", "邀鷙", "卋卋"]):
                continue
            lines.append(text)
    return merge_fragment_lines(lines)


def merge_fragment_lines(lines: list[str]) -> list[str]:
    merged: list[str] = []
    buffer = ""

    def flush() -> None:
        nonlocal buffer
        if buffer:
            merged.append(buffer.strip())
            buffer = ""

    for line in lines:
        is_boundary = (
            NUMBER_RE.match(line)
            or OPTION_RE.match(line)
            or line.startswith(("参考答案", "答案：", "答案:", "答：", "答:", "答案要点", "一、", "二、", "三、"))
            or "单选题" in line
            or "多选题" in line
            or "判断题" in line
            or "简答题" in line
        )
        if is_boundary:
            flush()
            merged.append(line)
        elif buffer:
            buffer += line
        else:
            buffer = line
    flush()
    return merged


def read_word(path: Path) -> list[str]:
    if path.suffix.lower() == ".docx":
        return split_embedded_markers(read_docx(path))

    raw = path.read_bytes().decode("utf-16le", errors="ignore")
    if "<body" in raw and "【" in raw:
        return split_embedded_markers(read_html_doc(path))
    return split_embedded_markers(read_binary_doc_runs(path))


def parse_answer_map(lines: list[str], single_mode: bool = False) -> dict[int, list[str]]:
    text = " ".join(clean_text(line) for line in lines)
    text = text.replace("✦", " ").replace("：", ":")
    answers: dict[int, list[str]] = {}

    for start, end, letters in RANGE_ANSWER_RE.findall(text):
        start_i, end_i = int(start), int(end)
        for offset, letter in enumerate(letters.upper()):
            number = start_i + offset
            if number <= end_i:
                answers[number] = [letter]

    if single_mode:
        for number, letters in re.findall(r"(\d+)\s*([A-F]{2,})", text, re.IGNORECASE):
            start_i = int(number)
            letters = letters.upper()
            # Some single-choice answer keys omit the range suffix, e.g. "16AACAB".
            if len(letters) > 1:
                for offset, letter in enumerate(letters):
                    answers[start_i + offset] = [letter]

    for number, letters in ANSWER_ITEM_RE.findall(text):
        if single_mode and len(letters) > 1:
            continue
        answers[int(number)] = list(letters.upper())

    # Some source files write "参考答案 ACD 2.BC..." and omit "1.".
    first = re.match(r".*?参考答案\s*([A-F]{1,6})(?:\s+|\s*\d)", text, re.IGNORECASE)
    if first and 1 not in answers:
        answers[1] = list(first.group(1).upper())

    orphans = re.findall(r"(?:^|\s)[\.、]\s*([A-F]{1,6})(?=\s+\d|$)", text, re.IGNORECASE)
    if orphans and answers:
        max_number = max(answers)
        for letters in orphans:
            missing = next((number for number in range(1, max_number + 1) if number not in answers), None)
            if missing is not None:
                answers[missing] = list(letters.upper())

    return answers


def section_bounds(lines: list[str], start_tokens: Iterable[str], end_tokens: Iterable[str]) -> tuple[int, int] | None:
    start = next((i for i, line in enumerate(lines) if any(token in line for token in start_tokens)), -1)
    if start < 0:
        return None
    end = len(lines)
    for i in range(start + 1, len(lines)):
        if any(token in lines[i] for token in end_tokens):
            end = i
            break
    return start + 1, end


def parse_choice_section(lines: list[str], qtype: str, chapter: str, answer_map: dict[int, list[str]]) -> list[ParsedQuestion]:
    questions: list[ParsedQuestion] = []
    current_number = 0
    current_question = ""
    current_options: list[str] = []

    def finish() -> None:
        nonlocal current_number, current_question, current_options
        if current_question and current_options:
            questions.append(
                ParsedQuestion(
                    type=qtype,
                    question=current_question,
                    options=current_options,
                    answer=answer_map.get(current_number, []),
                    analysis="",
                    chapter=chapter,
                )
            )
        current_number = 0
        current_question = ""
        current_options = []

    for line in lines:
        if line.startswith("参考答案"):
            break
        option_match = OPTION_RE.match(line)
        number_match = NUMBER_RE.match(line)

        if number_match:
            finish()
            current_number = int(number_match.group(1))
            current_question = clean_text(number_match.group(2))
            current_options = []
        elif option_match and current_question:
            option_text = clean_text(option_match.group(2))
            current_options.append(option_text or option_match.group(1).upper())
        elif current_question and len(current_options) < 2:
            current_question = clean_text(current_question + line)
        elif current_question and len(current_options) >= 2 and not option_match:
            # A few files lose the question number; after four options this is a new question.
            if len(current_options) >= 4:
                finish()
                current_number = len(questions) + 1
                current_question = line
            else:
                current_options[-1] = clean_text(current_options[-1] + line)
    finish()
    return questions


def collect_answer_lines(lines: list[str], start: int, end: int) -> list[str]:
    answer_start = -1
    for i in range(start, end):
        if lines[i].startswith("参考答案"):
            answer_start = i
            break
    if answer_start < 0:
        answer_start = next((i for i in range(start, end) if re.match(r"^\d+[\.、]?[A-F]+$", lines[i], re.IGNORECASE)), -1)
    if answer_start < 0:
        return []
    return lines[answer_start:end]


def parse_standard_document(lines: list[str], chapter: str) -> list[ParsedQuestion]:
    questions: list[ParsedQuestion] = []

    single_bounds = section_bounds(lines, ["单选题"], ["多选题", "判断题", "简答题", "二、简答题"])
    multi_bounds = section_bounds(lines, ["多选题"], ["判断题", "简答题", "二、简答题"])

    if single_bounds:
        start, end = single_bounds
        answers = parse_answer_map(collect_answer_lines(lines, start, end), single_mode=True)
        questions.extend(parse_choice_section(lines[start:end], "single", chapter, answers))

    if multi_bounds:
        start, end = multi_bounds
        answers = parse_answer_map(collect_answer_lines(lines, start, end))
        questions.extend(parse_choice_section(lines[start:end], "multiple", chapter, answers))

    essay_start = next((i for i, line in enumerate(lines) if "简答题" in line or "思考题" in line), -1)
    if essay_start >= 0:
        current_q = ""
        current_answer: list[str] = []
        for line in lines[essay_start + 1 :]:
            number_match = NUMBER_RE.match(line)
            answer_match = re.match(r"^(答|答案要点)[:：]?\s*(.*)$", line)
            if number_match:
                if current_q:
                    questions.append(ParsedQuestion("essay", current_q, [], ["\n".join(current_answer).strip()], "", chapter))
                current_q = clean_text(number_match.group(2))
                current_answer = []
            elif answer_match and current_q:
                current_answer.append(clean_text(answer_match.group(2)))
            elif current_q and current_answer:
                current_answer.append(line)
            elif current_q and not current_answer:
                current_q = clean_text(current_q + line)
        if current_q:
            questions.append(ParsedQuestion("essay", current_q, [], ["\n".join(current_answer).strip()], "", chapter))

    return questions


def parse_ai_html_document(lines: list[str], chapter: str) -> list[ParsedQuestion]:
    questions: list[ParsedQuestion] = []
    current: ParsedQuestion | None = None

    def finish() -> None:
        nonlocal current
        if current and current.question:
            questions.append(current)
        current = None

    for line in lines:
        type_match = AI_TYPE_RE.search(line)
        if type_match:
            finish()
            qtype = {"单选题": "single", "多选题": "multiple", "判断题": "judge"}[type_match.group(1)]
            question = clean_text(line[type_match.end() :])
            question = re.sub(r"^[#\d、\s]+", "", question)
            current = ParsedQuestion(qtype, question, [], [], "", chapter)
            continue

        if not current:
            continue

        option_match = OPTION_RE.match(line)
        if option_match and not line.startswith("答案"):
            option_text = clean_text(option_match.group(2))
            current.options.append(option_text)
        elif line.startswith(("答案：", "答案:")):
            answer_text = line.split("：", 1)[-1] if "：" in line else line.split(":", 1)[-1]
            if current.type == "judge":
                current.options = ["正确", "错误"]
                if any(word in answer_text for word in ["正确", "对", "是", "T", "true"]):
                    current.answer = ["正确"]
                elif any(word in answer_text for word in ["错误", "错", "否", "F", "false"]):
                    current.answer = ["错误"]
            else:
                letters = re.sub(r"[^A-Fa-f]", "", answer_text).upper()
                current.answer = list(letters)
        elif not current.options and not line.startswith(("答案", "解析")):
            current.question = clean_text(current.question + line)

    finish()
    return questions


def parse_file(path: Path) -> list[ParsedQuestion]:
    lines = read_word(path)
    chapter = path.stem.replace("《习思想概论》", "").replace("知识点及练习题", "").strip(" -_")
    if any("【单选题】" in line or "【多选题】" in line for line in lines):
        return parse_ai_html_document(lines, chapter or "人工智能基础与应用")
    return parse_standard_document(lines, chapter or "未分类")


def make_id(index: int) -> str:
    return f"imported-{index:04d}"


def main() -> None:
    if len(sys.argv) < 3:
        raise SystemExit("Usage: convert_word_banks.py OUTPUT.json INPUT...")

    output = Path(sys.argv[1])
    inputs = [Path(arg) for arg in sys.argv[2:]]
    all_questions: list[dict[str, object]] = []
    report: list[dict[str, object]] = []

    for path in inputs:
        if path.name.startswith("~$"):
            continue
        parsed = parse_file(path)
        report.append({"file": str(path), "count": len(parsed)})
        for item in parsed:
            all_questions.append(
                {
                    "id": make_id(len(all_questions) + 1),
                    "type": item.type,
                    "question": item.question,
                    "options": item.options,
                    "answer": item.answer,
                    "analysis": item.analysis,
                    "chapter": item.chapter,
                }
            )

    payload = {
        "title": "习题转换题库",
        "questions": all_questions,
        "meta": {
            "sourceCount": len(inputs),
            "convertedCount": len(all_questions),
            "report": report,
        },
    }
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(payload["meta"], ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
